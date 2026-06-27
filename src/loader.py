"""
Document loader for H-RAG pipeline.
Walks a folder, reads .md, .txt, .pdf, and .docx files,
and returns structured Document objects with metadata.
"""

from __future__ import annotations

import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

import fitz  # PyMuPDF for PDFs
from docx import Document as DocxDocument  # python-docx for DOCX

from src.config import SOURCE_PATH, PROJECTS_PATH


@dataclass
class Document:
    """A loaded document with content and metadata."""
    content: str
    metadata: dict = field(default_factory=dict)

    @property
    def doc_id(self) -> str:
        """Deterministic ID based on file path."""
        return hashlib.sha256(
            self.metadata.get("file_path", self.content[:100]).encode()
        ).hexdigest()[:16]


SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


def extract_text(file_path: Path) -> str:
    """Extract text from supported file types."""
    ext = file_path.suffix.lower()
    
    if ext in {".md", ".txt"}:
        return file_path.read_text(encoding="utf-8", errors="replace")
    
    elif ext == ".pdf":
        text_parts = []
        try:
            with fitz.open(str(file_path)) as doc:
                for page in doc:
                    text_parts.append(page.get_text())
            return "\n".join(text_parts)
        except Exception as e:
            print(f"Warning: Failed to parse PDF {file_path}: {e}")
            return ""
            
    elif ext == ".docx":
        try:
            doc = DocxDocument(str(file_path))
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            print(f"Warning: Failed to parse DOCX {file_path}: {e}")
            return ""
            
    return ""


def load_documents(source_path: Path | None = None, node_name: str = "anthropic") -> List[Document]:
    """
    Recursively load files from the source folder.

    Returns a list of Document objects, each with metadata.
    """
    root = source_path or SOURCE_PATH

    if not root.exists():
        raise FileNotFoundError(f"Source folder not found: {root}")

    documents: List[Document] = []

    for file_path in sorted(root.rglob("*")):
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        if not file_path.is_file():
            continue

        content = extract_text(file_path)
        if not content.strip():
            continue

        # Compute relative path from source root
        rel_path = file_path.relative_to(root)
        parent_rel = rel_path.parent

        # Folder name: immediate parent, or "root" if file is at top level
        folder_name = parent_rel.parts[-1] if parent_rel.parts else "root"

        metadata = {
            "node": node_name,
            "file_name": file_path.name,
            "file_path": str(rel_path),
            "file_type": file_path.suffix.lower(),
            "folder": folder_name,
            "folder_path": str(parent_rel) if parent_rel.parts else "root",
            "char_count": len(content),
        }

        documents.append(Document(content=content, metadata=metadata))

    return documents


def get_folder_groups(documents: List[Document]) -> dict[str, List[Document]]:
    """
    Group documents by their folder path.
    Returns a dict mapping folder_path -> list of Documents.
    """
    groups: dict[str, List[Document]] = {}
    for doc in documents:
        folder = doc.metadata["folder_path"]
        groups.setdefault(folder, []).append(doc)
    return groups
